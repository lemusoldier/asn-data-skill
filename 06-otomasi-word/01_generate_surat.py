#!/usr/bin/env python3
"""01 — Generate Surat Dinas Individu
Membuat satu file Word (.docx) secara otomatis.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def buat_surat_dinas(nama, nip, jabatan, instansi):
    doc = Document()
    
    # Set margin ke 2.54 cm (1 inch)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 1. KOP SURAT (Simulasi)
    p_kop = doc.add_paragraph()
    p_kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_kop1 = p_kop.add_run("PEMERINTAH KABUPATEN REMBANG\n")
    run_kop1.bold = True
    run_kop1.font.size = Pt(14)
    
    run_kop2 = p_kop.add_run("DINAS PENDIDIKAN DAN KEBUDAYAAN\n")
    run_kop2.bold = True
    run_kop2.font.size = Pt(16)
    
    run_kop3 = p_kop.add_run("Jl. Raya Rembang - Blora Km. 2, Rembang, Jawa Tengah\n")
    run_kop3.font.size = Pt(10)
    run_kop3.italic = True
    
    # Garis pembatas Kop
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.add_run("_" * 60).bold = True
    
    # 2. NOMOR & HAL
    doc.add_paragraph(f"Nomor : 800/1024/2026\n"
                      f"Sifat : Penting\n"
                      f"Lamp  : -\n"
                      f"Hal   : Pemanggilan Pembekalan Analisis Data")
    
    # Tanggal
    p_tgl = doc.add_paragraph(f"Rembang, {datetime.now().strftime('%d %B %Y')}")
    p_tgl.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Penerima
    doc.add_paragraph(f"Kepada Yth.\n"
                      f"Sdr. {nama}\n"
                      f"NIP. {nip}\n"
                      f"Jabatan: {jabatan}\n"
                      f"Di - {instansi}")
                      
    # 3. ISI SURAT
    doc.add_paragraph("Dengan hormat,\n"
                      "Menindaklanjuti program peningkatan kapasitas digital aparatur sipil negara di lingkungan Pemerintah Kabupaten Rembang, dengan ini kami mengharap kehadiran Saudara pada:")
                      
    # Detail Acara (Bullet/Indented)
    p_detail = doc.add_paragraph()
    p_detail.paragraph_format.left_indent = Inches(0.5)
    p_detail.add_run("Hari/Tanggal : Senin, 22 Juni 2026\n"
                     "Waktu        : 09.00 WIB - Selesai\n"
                     "Tempat       : Aula Gedung B, Dinas Pendidikan\n"
                     "Acara        : Workshop Pengolahan Data menggunakan Python dan Excel")
                     
    doc.add_paragraph("Demikian surat undangan ini kami sampaikan. Mengingat pentingnya acara tersebut, kehadiran Saudara sangat kami harapkan. Atas perhatian dan kerjasamanya diucapkan terima kasih.")
    
    # 4. TANDA TANGAN (TTE Simulasi)
    p_ttd = doc.add_paragraph()
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ttd.paragraph_format.space_before = Pt(30)
    p_ttd.add_run("Kepala Dinas Pendidikan\n\n\n\n"
                 "Drs. SUTRISNO, M.Pd\n"
                 "Pembina Utama Muda\n"
                 "NIP. 197005121995031002")
                 
    # Simpan
    nama_file = f"Surat_Undangan_{nama.replace(' ', '_')}.docx"
    doc.save(nama_file)
    print(f"✅ Berhasil membuat surat untuk {nama} -> {nama_file}")

# Test run dengan data pertama
buat_surat_dinas("Budi Santoso", "198501152010011004", "Analis Kepegawaian", "Dinas Pendidikan")
