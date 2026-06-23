#!/usr/bin/env python3
"""02 — Bulk Undangan Rapat Massal
Membaca file data_undangan.csv dan membuat surat .docx satu per satu otomatis.
"""
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Buat folder output
output_folder = "output_surat"
os.makedirs(output_folder, exist_ok=True)

# Baca CSV
df = pd.read_csv("data_undangan.csv")

print(f"Mulai membuat surat massal untuk {len(df)} orang...")
print("-" * 50)

for index, row in df.iterrows():
    doc = Document()
    
    # Margin
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Kop
    p_kop = doc.add_paragraph()
    p_kop.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_kop.add_run("PEMERINTAH KABUPATEN REMBANG\n")
    r1.bold = True
    r1.font.size = Pt(13)
    r2 = p_kop.add_run("BADAN KEPEGAWAIAN DAERAH (BKD)\n")
    r2.bold = True
    r2.font.size = Pt(15)
    r3 = p_kop.add_run("Jl. Pahlawan No. 45, Rembang\n")
    r3.italic = True
    r3.font.size = Pt(9)
    
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.add_run("_" * 65).bold = True
    
    # Isi
    doc.add_paragraph(f"Nomor : 800/{100+index}/2026\n"
                      f"Sifat : Penting\n"
                      f"Hal   : Undangan Ujian Kompetensi Digital")
                      
    doc.add_paragraph(f"Kepada Yth.\n"
                      f"Sdr/i. {row['nama']}\n"
                      f"Jabatan: {row['jabatan']}\n"
                      f"Di - {row['instansi']}")
                      
    doc.add_paragraph("Dengan hormat,\n"
                      "Diberitahukan kepada Saudara untuk hadir dalam kegiatan evaluasi kecakapan olah data yang diselenggarakan pada:")
                      
    p_detail = doc.add_paragraph()
    p_detail.paragraph_format.left_indent = Inches(0.5)
    p_detail.add_run("Hari/Tanggal : Rabu, 24 Juni 2026\n"
                     "Pukul        : 08.00 WIB s/d Selesai\n"
                     "Tempat       : Lab Komputer BKD Rembang\n"
                     "Keterangan   : Membawa laptop masing-masing.")
                     
    doc.add_paragraph("Kehadiran bersifat wajib dan tidak dapat diwakilkan. Terima kasih.")
    
    # TTD
    p_ttd = doc.add_paragraph()
    p_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ttd.paragraph_format.space_before = Pt(30)
    p_ttd.add_run("Kepala BKD Rembang\n\n\n\n"
                 "ARIEF ROHMAN, S.Sos\n"
                 "NIP. 197508192003121004")
                 
    # Simpan ke folder output
    file_path = os.path.join(output_folder, f"Undangan_BKD_{row['nama'].replace(' ', '_')}.docx")
    doc.save(file_path)
    print(f"[{index+1}/{len(df)}] ✅ Selesai: {file_path}")

print("-" * 50)
print(f"Sukses! Semua surat disimpan di folder: {output_folder}/")
